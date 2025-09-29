from azure.storage.blob.aio import BlobServiceClient
from azure.storage.blob import ContentSettings
import uuid
import os
import io
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup

from dotenv import load_dotenv
from fastapi import HTTPException, UploadFile, status

load_dotenv()

storage_account_connection_string = os.getenv("STORAGE_ACCOUNT_CONNECTION_STRING")
container_name = "assignment-files"

async def upload_to_blob_old(file: UploadFile, file_bytes: bytes):
    blob_service_client = BlobServiceClient.from_connection_string(storage_account_connection_string)
    blob_name = f"{uuid.uuid4()}_{file.filename}"
    blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_name)

    await blob_client.upload_blob(file_bytes, overwrite=True,
        content_settings=ContentSettings(content_type=file.content_type)
    )
    return blob_client.url


async def upload_to_blob(file: UploadFile, file_bytes: bytes):
    try:
        async with BlobServiceClient.from_connection_string(storage_account_connection_string) as blob_service_client:
            blob_name = f"{uuid.uuid4()}_{file.filename}"
            async with blob_service_client.get_blob_client(container=container_name, blob=blob_name) as blob_client:
                await blob_client.upload_blob(file_bytes, overwrite=True,
                    content_settings=ContentSettings(content_type=file.content_type)
                )
                return blob_client.url
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to upload file to Azure Blob Storage: {str(e)}"
        )


def html_to_word_document(html_content: str, title: str) -> Document:
    """Convert HTML content to a Word document with basic formatting."""
    doc = Document()

    # Add title as heading
    doc.add_heading(title, 0)

    # Parse HTML content
    soup = BeautifulSoup(html_content, 'html.parser')

    # Process each element
    for element in soup:
        if element.name == 'p':
            # Add paragraph
            paragraph = doc.add_paragraph()
            add_text_with_formatting(paragraph, element)
        elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Add heading
            level = int(element.name[1])
            doc.add_heading(element.get_text(strip=True), level)
        elif element.name == 'ul':
            # Add unordered list
            for li in element.find_all('li'):
                p = doc.add_paragraph(li.get_text(strip=True), style='List Bullet')
        elif element.name == 'ol':
            # Add ordered list
            for li in element.find_all('li'):
                p = doc.add_paragraph(li.get_text(strip=True), style='List Number')
        elif element.name is None and element.strip():
            # Plain text node
            doc.add_paragraph(element.strip())

    return doc


def add_text_with_formatting(paragraph, element):
    """Add text to paragraph with basic formatting (bold, italic)."""
    for content in element.children:
        if content.name == 'strong' or content.name == 'b':
            run = paragraph.add_run(content.get_text())
            run.bold = True
        elif content.name == 'em' or content.name == 'i':
            run = paragraph.add_run(content.get_text())
            run.italic = True
        elif content.name is None:
            # Plain text
            paragraph.add_run(str(content))
        else:
            # Other tags, just get text
            paragraph.add_run(content.get_text())


async def upload_html_as_word_to_blob(html_content: str, title: str, student_id: int) -> str:
    """Convert HTML content to Word document and upload to blob storage."""
    try:
        # Generate Word document
        doc = html_to_word_document(html_content, title)

        # Save document to memory
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        # Create blob name
        safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_title = safe_title.replace(' ', '_')
        blob_name = f"{student_id}_{safe_title}_{uuid.uuid4().hex}.docx"

        # Upload to blob storage
        async with BlobServiceClient.from_connection_string(storage_account_connection_string) as blob_service_client:
            async with blob_service_client.get_blob_client(container=container_name, blob=blob_name) as blob_client:
                await blob_client.upload_blob(
                    doc_bytes.getvalue(),
                    overwrite=True,
                    content_settings=ContentSettings(
                        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                )
                return blob_client.url

    except Exception as e:
        print(f"Failed to create/upload Word document: {str(e)}")
        # Return None to indicate failure, let calling code handle gracefully
        return None