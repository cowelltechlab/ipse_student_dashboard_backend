from fastapi import HTTPException
from datetime import datetime
from application.features.versionHistory.schemas import AssignmentVersionResponse, AssignmentVersionUpdate
from azure.cosmos import exceptions
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
import io
import re

# Import legacy conversion utilities
from application.features.assignment_version_generation.crud import convert_json_to_html


def get_html_content_from_version_document(document: dict) -> str:
    """
    Extract HTML content from version document, handling both new and legacy formats.

    Args:
        document: Version document from database

    Returns:
        HTML content string
    """
    final_content = document.get("final_generated_content", {})

    # Check for new HTML format first
    if "html_content" in final_content:
        return final_content["html_content"]

    # Check for legacy JSON format
    if "json_content" in final_content:
        json_content = final_content["json_content"]
        return convert_json_to_html(json_content)

    # Check for very old format with direct HTML fields
    if "generated_html" in final_content:
        return final_content["generated_html"]

    if "raw_text" in final_content:
        return final_content["raw_text"]

    # Fallback - no content found
    return "<p>No content available</p>"


def convert_html_to_word_bytes(html_content: str) -> bytes:
    """
    Convert HTML content to Word document bytes
    """
    doc = Document()
    
    if not html_content:
        doc.add_paragraph("No content available")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Track numbered lists to ensure each section starts fresh
    numbered_list_count = 0
    
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'pre', 'div', 'hr']):
        if element.name in ['h1', 'h2', 'h3']:
            # Add headings - reset numbered list counter when we hit a new section
            level = int(element.name[1])
            doc.add_heading(element.get_text(strip=True), level=level)
            numbered_list_count = 0  # Reset counter for new sections
        
        elif element.name == 'hr':
            # Add a horizontal line separator instead of page break
            paragraph = doc.add_paragraph()
            # paragraph.add_run("_" * 50)  # Add a line of underscores as separator
            numbered_list_count = 0  # Reset counter at section breaks
        
        elif element.name == 'p':
            # Add paragraphs
            text = element.get_text(strip=True)
            if text:
                paragraph = doc.add_paragraph(text)
        
        elif element.name in ['ul', 'ol']:
            # Handle lists - ensure numbered lists restart properly
            is_numbered = element.name == 'ol'
            
            if is_numbered:
                numbered_list_count += 1
                # For numbered lists, create a new numbering instance to ensure fresh start
                for i, li in enumerate(element.find_all('li')):
                    text = li.get_text(strip=True)
                    if text:
                        paragraph = doc.add_paragraph(text, style='List Number')
                        # Restart numbering for each new <ol> section
                        if i == 0:  # First item in this list
                            try:
                                paragraph._element.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = numbered_list_count
                            except:
                                # Fallback if numbering manipulation fails
                                pass
            else:
                # Unordered lists
                for li in element.find_all('li'):
                    text = li.get_text(strip=True)
                    if text:
                        doc.add_paragraph(text, style='List Bullet')
        
        elif element.name == 'pre':
            # Handle code blocks
            text = element.get_text()
            if text:
                paragraph = doc.add_paragraph(text)
                paragraph.style = 'Normal'
                # Make code blocks appear in monospace-like formatting
                for run in paragraph.runs:
                    run.font.name = 'Courier New'
        
        elif element.name == 'div':
            # Handle different div types
            if 'ql-code-block' in element.get('class', []):
                # Handle Quill code blocks
                text = element.get_text()
                if text.strip():
                    paragraph = doc.add_paragraph(text)
                    for run in paragraph.runs:
                        run.font.name = 'Courier New'
            elif 'counter-reset' in element.get('style', ''):
                # This div marks a section boundary - reset numbering
                numbered_list_count = 0
                # Process child elements
                for child in element.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'pre']):
                    # Re-process child elements within this div
                    pass
    
    # Convert to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()



def get_assignment_version_by_doc_id(container, document_version_id: str) -> AssignmentVersionResponse:
    query = "SELECT * FROM c WHERE c.id = @id"
    parameters = [{"name": "@id", "value": document_version_id}]

    try:
        items = list(container.query_items(
            query=query,
            parameters=parameters,
            enable_cross_partition_query=True  # keep this if needed
        ))
        if not items:
            raise HTTPException(status_code=404, detail="Version not found")
        return AssignmentVersionResponse(**items[0])
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch version: {str(e)}")
    

def download_assignment_version(container, document_version_id: str) -> dict:
    query = "SELECT * FROM c WHERE c.id = @id"
    parameters = [{"name": "@id", "value": document_version_id}]

    try:
        items = list(container.query_items(
            query=query,
            parameters=parameters,
            enable_cross_partition_query=True
        ))
        if not items:
            raise HTTPException(status_code=404, detail="Version not found")
        
        item = items[0]
        
        # Get HTML content using the unified helper that handles all formats
        combined_html = get_html_content_from_version_document(item)

        if not combined_html or combined_html == "<p>No content available</p>":
            raise HTTPException(status_code=400, detail="No final content available for download")
        
        # Convert to Word document
        word_bytes = convert_html_to_word_bytes(combined_html)
        
        # Generate a descriptive filename
        assignment_id = item.get("assignment_id", "unknown")
        version_number = item.get("version_number", "1")
        filename = f"assignment_{assignment_id}_v{version_number}.docx"

        return {
            "file_name": filename,
            "file_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "file_content": word_bytes
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to download version: {str(e)}")



def delete_version_by_assignment_version(container, assignment_id: str, version_number: int):
    try:
        # 1. Find document by assignment_id and version_number
        query = """
        SELECT * FROM c 
        WHERE c.assignment_id = @assignment_id AND c.version_number = @version_number
        """
        params = [
            {"name": "@assignment_id", "value": assignment_id},
            {"name": "@version_number", "value": version_number}
        ]
        items = list(container.query_items(query=query, parameters=params, enable_cross_partition_query=True))

        if not items:
            raise HTTPException(status_code=404, detail="Document not found")

        item = items[0]
        doc_id = item["id"]
        modifier_id = item["modifier_id"]

        #2.  Delete using partition key
        container.delete_item(item=doc_id, partition_key=str(modifier_id))

    except exceptions.CosmosResourceNotFoundError:
        raise HTTPException(status_code=404, detail="Document not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to delete version: {str(e)}")

def update_version(container, assignment_id: str, version_number: int, update_data: AssignmentVersionUpdate) -> AssignmentVersionResponse:
    # 1. Find existing version
    query = """
    SELECT * FROM c
    WHERE c.assignment_id = @assignment_id AND c.version_number = @version_number
    """
    params = [
        {"name": "@assignment_id", "value": assignment_id},
        {"name": "@version_number", "value": version_number}
    ]
    items = list(container.query_items(query=query, parameters=params, enable_cross_partition_query=True))
    
    if not items:
        raise HTTPException(status_code=404, detail="Version not found")

    existing = items[0]
    doc_id = existing["id"]
    partition_key = existing["modifier_id"]

    # 2. Prepare the update
    update_dict = update_data.dict(exclude_unset=True)

    # Convert date_modified if provided
    if "date_modified" in update_dict and isinstance(update_dict["date_modified"], datetime):
        update_dict["date_modified"] = update_dict["date_modified"].isoformat()

    # Apply updates to existing doc
    for k, v in update_dict.items():
        existing[k] = v

    try:
        # If this update sets finalized=True, unset others
        if update_dict.get("finalized") is True:
            # Find all versions for this assignment
            query_all = "SELECT * FROM c WHERE c.assignment_id = @assignment_id"
            params_all = [{"name": "@assignment_id", "value": assignment_id}]
            all_versions = list(container.query_items(
                query=query_all,
                parameters=params_all,
                enable_cross_partition_query=True
            ))

            for v in all_versions:
                if v["id"] != doc_id and v.get("finalized"):
                    v["finalized"] = False
                    container.replace_item(item=v["id"], body=v)

        # Save the current version
        container.replace_item(item=doc_id, body=existing)
        existing["modifier_id"] = int(existing["modifier_id"])
        return AssignmentVersionResponse(**existing)

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to update version: {str(e)}")
    
def finalize_by_id(container, assignment_version_id: str, finalized: bool) -> AssignmentVersionResponse:
    # Step 1: Get the current document by ID
    try:
        items = list(container.query_items(
            query="SELECT * FROM c WHERE c.id = @id",
            parameters=[{"name": "@id", "value": assignment_version_id}],
            enable_cross_partition_query=True
        ))

        if not items:
            raise HTTPException(status_code=404, detail="Assignment version not found")

        current = items[0]
        assignment_id = current["assignment_id"]
        current["finalized"] = finalized
        container.replace_item(item=current["id"], body=current)

        # Step 2: If setting to True, un-finalize all others
        if finalized:
            others = list(container.query_items(
                query="""
                SELECT * FROM c 
                WHERE c.assignment_id = @assignment_id AND c.id != @id
                """,
                parameters=[
                    {"name": "@assignment_id", "value": assignment_id},
                    {"name": "@id", "value": assignment_version_id}
                ],
                enable_cross_partition_query=True
            ))

            for item in others:
                if item.get("finalized") is True:
                    item["finalized"] = False
                    container.replace_item(item=item["id"], body=item)

        # Return updated
        current["modifier_id"] = int(current["modifier_id"])
        return AssignmentVersionResponse(**current)

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to finalize version: {str(e)}")
