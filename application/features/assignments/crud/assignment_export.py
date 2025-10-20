"""
Assignment export functionality - JSON and ZIP downloads
"""
from fastapi import HTTPException
import pyodbc
from application.database.mssql_connection import get_sql_db_connection
from application.database.nosql_connection import get_container
from datetime import datetime
from typing import List, Dict, Optional
import zipfile
import io
import csv

from application.features.versionHistory.crud import get_html_content_from_version_document, convert_html_to_word_bytes
from application.features.student_profile.crud import get_complete_profile


def export_student_assignments_json(student_id: int, assignment_ids: Optional[List[int]] = None) -> dict:
    """
    Export all assignment data for a student in JSON format.
    Includes student info, class associations, and all assignment data with versions and ratings.

    Args:
        student_id: The student's internal ID
        assignment_ids: Optional list of assignment IDs to filter by

    Returns:
        Dictionary with student data, classes, and assignments
    """
    try:
        with get_sql_db_connection() as conn:
            cursor = conn.cursor()

            # 1. Get student info
            cursor.execute("""
                SELECT
                    s.id, s.user_id, s.reading_level, s.writing_level, s.group_type,
                    u.first_name, u.last_name, u.email, u.gt_email,
                    y.name AS year_name
                FROM Students s
                INNER JOIN Users u ON s.user_id = u.id
                LEFT JOIN Years y ON s.year_id = y.id
                WHERE s.id = ?
            """, (student_id,))

            student_row = cursor.fetchone()
            if not student_row:
                raise HTTPException(status_code=404, detail=f"Student with id {student_id} not found")

            student_columns = [column[0] for column in cursor.description]
            student_data = dict(zip(student_columns, student_row))

            # 2. Get class associations with learning goals
            cursor.execute("""
                SELECT
                    c.id AS class_id, c.name AS class_name, c.course_code, c.term, c.type,
                    sc.learning_goal
                FROM StudentClasses sc
                INNER JOIN Classes c ON sc.class_id = c.id
                WHERE sc.student_id = ?
            """, (student_id,))

            class_rows = cursor.fetchall()
            class_columns = [column[0] for column in cursor.description]
            classes_data = [dict(zip(class_columns, row)) for row in class_rows]

            # 3. Get assignments (optionally filtered)
            if assignment_ids:
                placeholders = ",".join("?" for _ in assignment_ids)
                assignments_query = f"""
                    SELECT
                        a.id AS assignment_id, a.title, a.content, a.html_content,
                        a.date_created, a.blob_url, a.source_format, a.assignment_type_id,
                        c.id AS class_id, c.name AS class_name, c.course_code,
                        at.type AS assignment_type
                    FROM Assignments a
                    LEFT JOIN Classes c ON a.class_id = c.id
                    LEFT JOIN AssignmentTypes at ON a.assignment_type_id = at.id
                    WHERE a.student_id = ? AND a.id IN ({placeholders})
                    ORDER BY a.date_created DESC
                """
                cursor.execute(assignments_query, (student_id, *assignment_ids))
            else:
                assignments_query = """
                    SELECT
                        a.id AS assignment_id, a.title, a.content, a.html_content,
                        a.date_created, a.blob_url, a.source_format, a.assignment_type_id,
                        c.id AS class_id, c.name AS class_name, c.course_code,
                        at.type AS assignment_type
                    FROM Assignments a
                    LEFT JOIN Classes c ON a.class_id = c.id
                    LEFT JOIN AssignmentTypes at ON a.assignment_type_id = at.id
                    WHERE a.student_id = ?
                    ORDER BY a.date_created DESC
                """
                cursor.execute(assignments_query, (student_id,))

            assignment_rows = cursor.fetchall()
            assignment_columns = [column[0] for column in cursor.description]
            assignments_data = []

            # 4. For each assignment, get versions from Cosmos
            container = get_container()

            for row in assignment_rows:
                assignment = dict(zip(assignment_columns, row))

                # Query Cosmos for all versions of this assignment
                assignment_id = assignment["assignment_id"]
                query = f"SELECT * FROM c WHERE c.assignment_id = {assignment_id} ORDER BY c.version_number"
                versions = list(container.query_items(query=query, enable_cross_partition_query=True))

                # Structure class_info
                class_info = None
                if assignment.get("class_id"):
                    class_info = {
                        "id": assignment["class_id"],
                        "name": assignment["class_name"],
                        "course_code": assignment["course_code"]
                    }

                # Build assignment export object
                assignment_export = {
                    "assignment_id": assignment["assignment_id"],
                    "title": assignment["title"],
                    "content": assignment["content"],
                    "html_content": assignment.get("html_content"),
                    "date_created": assignment["date_created"].isoformat() if assignment["date_created"] else None,
                    "blob_url": assignment.get("blob_url"),
                    "source_format": assignment.get("source_format"),
                    "assignment_type": assignment.get("assignment_type"),
                    "assignment_type_id": assignment.get("assignment_type_id"),
                    "class_info": class_info,
                    "versions": versions
                }

                assignments_data.append(assignment_export)

        # 5. Build final response
        from datetime import datetime as dt
        export_data = {
            "student": student_data,
            "classes": classes_data,
            "assignments": assignments_data,
            "export_metadata": {
                "exported_at": dt.utcnow().isoformat(),
                "total_assignments": len(assignments_data),
                "filtered_by_ids": assignment_ids if assignment_ids else None
            }
        }

        return export_data

    except HTTPException:
        raise
    except pyodbc.Error as e:
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export error: {str(e)}")


def _format_student_info(student_data: dict) -> str:
    """Format student information as readable text"""
    lines = [
        "=" * 60,
        "STUDENT INFORMATION",
        "=" * 60,
        f"Name: {student_data.get('first_name', '')} {student_data.get('last_name', '')}",
        f"Student ID: {student_data.get('id', 'N/A')}",
        f"User ID: {student_data.get('user_id', 'N/A')}",
        f"Email: {student_data.get('email', 'N/A')}",
        f"GT Email: {student_data.get('gt_email', 'N/A')}",
        f"Year: {student_data.get('year_name', 'N/A')}",
        f"Reading Level: {student_data.get('reading_level', 'N/A')}",
        f"Writing Level: {student_data.get('writing_level', 'N/A')}",
        f"Group Type: {student_data.get('group_type', 'N/A')}",
        "=" * 60
    ]
    return "\n".join(lines)


def _format_classes_and_goals(classes_data: List[dict]) -> str:
    """Format class associations and learning goals as readable text"""
    lines = [
        "=" * 60,
        "CLASS ASSOCIATIONS & LEARNING GOALS",
        "=" * 60,
        ""
    ]

    if not classes_data:
        lines.append("No class associations found.")
    else:
        for i, cls in enumerate(classes_data, 1):
            lines.extend([
                f"Class {i}:",
                f"  Name: {cls.get('class_name', 'N/A')}",
                f"  Course Code: {cls.get('course_code', 'N/A')}",
                f"  Term: {cls.get('term', 'N/A')}",
                f"  Type: {cls.get('type', 'N/A')}",
                f"  Learning Goal: {cls.get('learning_goal', 'N/A')}",
                ""
            ])

    lines.append("=" * 60)
    return "\n".join(lines)


def _format_rating_data(rating_data: dict) -> str:
    """Format rating data as readable text"""
    if not rating_data:
        return "No rating data available.\n"

    lines = ["RATING DATA", "-" * 40, ""]

    def format_section(title: str, data: dict, indent: int = 0):
        result = []
        prefix = "  " * indent
        result.append(f"{prefix}{title}:")
        for key, value in data.items():
            if isinstance(value, dict):
                result.extend(format_section(key, value, indent + 1))
            elif isinstance(value, list):
                result.append(f"{prefix}  {key}: {', '.join(str(v) for v in value)}")
            else:
                result.append(f"{prefix}  {key}: {value}")
        return result

    for section_name, section_data in rating_data.items():
        if isinstance(section_data, dict):
            lines.extend(format_section(section_name, section_data))
            lines.append("")
        else:
            lines.append(f"{section_name}: {section_data}")

    return "\n".join(lines)


def _format_generated_options(generated_options: list) -> str:
    """Format generated options (learning pathways) as readable text"""
    if not generated_options:
        return "No learning pathways generated.\n"

    lines = ["LEARNING PATHWAYS (GENERATED OPTIONS)", "=" * 60, ""]

    for i, option in enumerate(generated_options, 1):
        selected_marker = " ✓ SELECTED" if option.get('selected') else ""
        lines.extend([
            f"Option {i}: {option.get('name', 'N/A')}{selected_marker}",
            "-" * 60,
            f"Internal ID: {option.get('internal_id', 'N/A')}",
            "",
            f"Description:",
            f"  {option.get('description', 'N/A')}",
            "",
            f"Why Good (Existing Skills):",
            f"  {option.get('why_good_existing', 'N/A')}",
            "",
            f"Why Challenge:",
            f"  {option.get('why_challenge', 'N/A')}",
            "",
            f"Why Good (Growth):",
            f"  {option.get('why_good_growth', 'N/A')}",
            "",
            f"Selection Logic:",
            f"  {option.get('selection_logic', 'N/A')}",
            "",
            ""
        ])

    return "\n".join(lines)


def _format_rating_history(rating_history: list) -> str:
    """Format rating history as readable text"""
    if not rating_history:
        return "No rating history available.\n"

    lines = ["RATING HISTORY", "=" * 60, ""]

    for i, entry in enumerate(rating_history, 1):
        timestamp = entry.get('timestamp', 'N/A')
        update_type = entry.get('update_type', 'N/A')
        lines.extend([
            f"Rating Update #{i}",
            "-" * 60,
            f"Timestamp: {timestamp}",
            f"Update Type: {update_type}",
            ""
        ])

        # Format the rating data within this history entry
        if entry.get('rating_data'):
            rating_text = _format_rating_data(entry['rating_data'])
            # Indent the rating data
            for line in rating_text.split('\n'):
                if line.strip():
                    lines.append(f"  {line}")

        lines.extend(["", ""])

    return "\n".join(lines)


def _format_generation_history(generation_history: list) -> str:
    """Format generation history as readable text"""
    if not generation_history:
        return "No generation history available.\n"

    lines = ["GENERATION HISTORY", "=" * 60, ""]

    for i, entry in enumerate(generation_history, 1):
        timestamp = entry.get('timestamp', 'N/A')
        generation_type = entry.get('generation_type', 'N/A')
        html_length = len(entry.get('html_content', ''))

        lines.extend([
            f"Generation #{i}",
            "-" * 60,
            f"Timestamp: {timestamp}",
            f"Generation Type: {generation_type}",
            f"Content Length: {html_length} characters",
            "",
            ""
        ])

    return "\n".join(lines)


def _format_complete_version_details(version: dict) -> str:
    """Format complete version details including all metadata, options, and history"""
    lines = [
        "=" * 80,
        "COMPLETE VERSION DETAILS",
        "=" * 80,
        "",
        "=== BASIC INFORMATION ===",
        f"Version Number: {version.get('version_number', 'N/A')}",
        f"Modifier ID: {version.get('modifier_id', 'N/A')}",
        f"Student ID: {version.get('student_id', 'N/A')}",
        f"Assignment ID: {version.get('assignment_id', 'N/A')}",
        f"Date Modified: {version.get('date_modified', 'N/A')}",
        f"Finalized: {'Yes' if version.get('finalized') else 'No'}",
        "",
        "=== SKILLS FOR SUCCESS ===",
        version.get('skills_for_success', 'N/A'),
        "",
        "=== SELECTED OPTIONS ===",
        ', '.join(version.get('selected_options', [])) if version.get('selected_options') else 'None',
        "",
        "=== STUDENT'S IDEAS BOX (Additional Edit Suggestions) ===",
        version.get('additional_edit_suggestions', 'N/A') or 'None provided',
        "",
        "",
    ]

    # Add learning pathways
    lines.append(_format_generated_options(version.get('generated_options', [])))
    lines.append("")

    # Add generation history
    lines.append(_format_generation_history(version.get('generation_history', [])))
    lines.append("")

    # Add rating history
    lines.append(_format_rating_history(version.get('rating_history', [])))
    lines.append("")

    lines.append("=" * 80)
    return "\n".join(lines)


def _create_assignments_summary_csv(assignments_data: List[dict]) -> str:
    """Create CSV summary of all assignments"""
    output = io.StringIO()
    writer = csv.writer(output)

    # Header
    writer.writerow([
        "Assignment ID",
        "Title",
        "Class",
        "Assignment Type",
        "Date Created",
        "Number of Versions",
        "Finalized",
        "Rating Status"
    ])

    # Data rows
    for assignment in assignments_data:
        versions = assignment.get("versions", [])
        finalized = any(v.get("finalized") for v in versions)
        has_ratings = any(v.get("rating_data") for v in versions)

        rating_status = "Pending"
        if has_ratings:
            if finalized and any(v.get("finalized") and v.get("rating_data") for v in versions):
                rating_status = "Rated"
            else:
                rating_status = "Partially Rated"

        writer.writerow([
            assignment.get("assignment_id", ""),
            assignment.get("title", ""),
            assignment.get("class_info", {}).get("name", "N/A") if assignment.get("class_info") else "N/A",
            assignment.get("assignment_type", "N/A"),
            assignment.get("date_created", ""),
            len(versions),
            "Yes" if finalized else "No",
            rating_status
        ])

    return output.getvalue()


def export_student_assignments_download(student_id: int, assignment_ids: Optional[List[int]] = None) -> bytes:
    """
    Export all assignment data for a student as a user-friendly ZIP file.

    ZIP structure:
    - student_info.txt
    - classes_and_goals.txt
    - assignments_summary.csv
    - assignments/
        - assignment_{id}_{title}/
            - original_assignment.docx
            - version_{n}.docx (generated content)
            - version_{n}_complete_details.txt (all metadata, learning pathways, skills, history)
            - ratings.txt

    Args:
        student_id: The student's internal ID
        assignment_ids: Optional list of assignment IDs to filter by

    Returns:
        ZIP file as bytes
    """
    try:
        # Get JSON export data first
        export_data = export_student_assignments_json(student_id, assignment_ids)

        # Create ZIP in memory
        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # 1. Add student_info.txt
            student_info_text = _format_student_info(export_data["student"])
            zip_file.writestr("student_info.txt", student_info_text)

            # 2. Add classes_and_goals.txt
            classes_text = _format_classes_and_goals(export_data["classes"])
            zip_file.writestr("classes_and_goals.txt", classes_text)

            # 3. Add assignments_summary.csv
            summary_csv = _create_assignments_summary_csv(export_data["assignments"])
            zip_file.writestr("assignments_summary.csv", summary_csv)

            # 4. Add individual assignment folders
            for assignment in export_data["assignments"]:
                assignment_id = assignment["assignment_id"]
                title = assignment["title"]
                # Sanitize title for folder name
                safe_title = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in title)[:50]
                folder_name = f"assignments/assignment_{assignment_id}_{safe_title}"

                # 4a. Add original assignment content as Word doc
                if assignment.get("html_content"):
                    original_word_bytes = convert_html_to_word_bytes(assignment["html_content"])
                    zip_file.writestr(f"{folder_name}/original_assignment.docx", original_word_bytes)
                elif assignment.get("content"):
                    # If no HTML, create simple text doc
                    from docx import Document
                    doc = Document()
                    doc.add_heading(assignment["title"], 0)
                    doc.add_paragraph(assignment["content"])
                    doc_buffer = io.BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    zip_file.writestr(f"{folder_name}/original_assignment.docx", doc_buffer.getvalue())

                # 4b. Add each version as separate Word doc and detailed metadata
                versions = assignment.get("versions", [])
                all_ratings = []

                for version in versions:
                    version_num = version.get("version_number", "unknown")
                    finalized_tag = "_finalized" if version.get("finalized") else ""

                    # Get HTML content from version
                    html_content = get_html_content_from_version_document(version)

                    # Convert HTML content to Word document
                    version_word_bytes = convert_html_to_word_bytes(html_content)

                    zip_file.writestr(
                        f"{folder_name}/version_{version_num}{finalized_tag}.docx",
                        version_word_bytes
                    )

                    # Add complete version details as text file
                    complete_details = _format_complete_version_details(version)
                    zip_file.writestr(
                        f"{folder_name}/version_{version_num}{finalized_tag}_complete_details.txt",
                        complete_details
                    )

                    # Collect ratings
                    if version.get("rating_data"):
                        all_ratings.append({
                            "version_number": version_num,
                            "finalized": version.get("finalized"),
                            "rating_data": version.get("rating_data")
                        })

                # 4c. Add ratings.txt if any ratings exist
                if all_ratings:
                    ratings_text = f"RATINGS FOR ASSIGNMENT {assignment_id}: {title}\n"
                    ratings_text += "=" * 60 + "\n\n"

                    for rating_info in all_ratings:
                        ratings_text += f"Version {rating_info['version_number']}"
                        if rating_info['finalized']:
                            ratings_text += " (FINALIZED)"
                        ratings_text += "\n" + "-" * 60 + "\n"
                        ratings_text += _format_rating_data(rating_info['rating_data'])
                        ratings_text += "\n\n"

                    zip_file.writestr(f"{folder_name}/ratings.txt", ratings_text)
                else:
                    zip_file.writestr(f"{folder_name}/ratings.txt", "No ratings available for this assignment.\n")

        zip_buffer.seek(0)
        return zip_buffer.getvalue()

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Download export error: {str(e)}")


def export_complete_student_data(student_id: int, assignment_ids: Optional[List[int]] = None) -> bytes:
    """
    Export complete student data combining profile and assignments in one comprehensive ZIP.

    Includes:
    - Full student profile (strengths, challenges, goals, interests)
    - All classes with learning goals
    - PowerPoint achievements URLs
    - All assignments with versions and ratings (same as assignment export)

    Args:
        student_id: The student's internal ID
        assignment_ids: Optional list of assignment IDs to filter by

    Returns:
        ZIP file as bytes with complete student data
    """
    try:
        # 1. Get student profile data
        try:
            profile_data = get_complete_profile(student_id)
            if not profile_data:
                raise HTTPException(status_code=404, detail=f"Student profile for id {student_id} not found")
        except Exception as e:
            raise HTTPException(status_code=404, detail=f"Failed to fetch student profile: {str(e)}")

        # 2. Get assignment export data
        assignment_export_data = export_student_assignments_json(student_id, assignment_ids)

        # 3. Create comprehensive ZIP
        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # === STUDENT PROFILE SECTION ===

            # Add detailed student profile
            profile_text = _format_complete_student_profile(profile_data)
            zip_file.writestr("student_profile.txt", profile_text)

            # Add profile as JSON for programmatic access
            import json
            zip_file.writestr("student_profile.json", json.dumps(profile_data, indent=2, default=str))

            # === CLASSES AND GOALS SECTION ===
            classes_text = _format_classes_and_goals(assignment_export_data["classes"])
            zip_file.writestr("classes_and_learning_goals.txt", classes_text)

            # === ASSIGNMENTS SECTION ===
            # Add assignments summary CSV
            summary_csv = _create_assignments_summary_csv(assignment_export_data["assignments"])
            zip_file.writestr("assignments_summary.csv", summary_csv)

            # Add individual assignment folders (reuse existing logic)
            for assignment in assignment_export_data["assignments"]:
                assignment_id = assignment["assignment_id"]
                title = assignment["title"]
                safe_title = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in title)[:50]
                folder_name = f"assignments/assignment_{assignment_id}_{safe_title}"

                # Original assignment
                if assignment.get("html_content"):
                    original_word_bytes = convert_html_to_word_bytes(assignment["html_content"])
                    zip_file.writestr(f"{folder_name}/original_assignment.docx", original_word_bytes)
                elif assignment.get("content"):
                    from docx import Document
                    doc = Document()
                    doc.add_heading(assignment["title"], 0)
                    doc.add_paragraph(assignment["content"])
                    doc_buffer = io.BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    zip_file.writestr(f"{folder_name}/original_assignment.docx", doc_buffer.getvalue())

                # Assignment versions
                versions = assignment.get("versions", [])
                all_ratings = []

                for version in versions:
                    version_num = version.get("version_number", "unknown")
                    finalized_tag = "_finalized" if version.get("finalized") else ""

                    html_content = get_html_content_from_version_document(version)
                    version_word_bytes = convert_html_to_word_bytes(html_content)

                    zip_file.writestr(
                        f"{folder_name}/version_{version_num}{finalized_tag}.docx",
                        version_word_bytes
                    )

                    # Add complete version details as text file
                    complete_details = _format_complete_version_details(version)
                    zip_file.writestr(
                        f"{folder_name}/version_{version_num}{finalized_tag}_complete_details.txt",
                        complete_details
                    )

                    if version.get("rating_data"):
                        all_ratings.append({
                            "version_number": version_num,
                            "finalized": version.get("finalized"),
                            "rating_data": version.get("rating_data")
                        })

                # Ratings file
                if all_ratings:
                    ratings_text = f"RATINGS FOR ASSIGNMENT {assignment_id}: {title}\n"
                    ratings_text += "=" * 60 + "\n\n"

                    for rating_info in all_ratings:
                        ratings_text += f"Version {rating_info['version_number']}"
                        if rating_info['finalized']:
                            ratings_text += " (FINALIZED)"
                        ratings_text += "\n" + "-" * 60 + "\n"
                        ratings_text += _format_rating_data(rating_info['rating_data'])
                        ratings_text += "\n\n"

                    zip_file.writestr(f"{folder_name}/ratings.txt", ratings_text)
                else:
                    zip_file.writestr(f"{folder_name}/ratings.txt", "No ratings available for this assignment.\n")

            # === EXPORT METADATA ===
            metadata_text = _format_export_metadata(assignment_export_data, profile_data)
            zip_file.writestr("export_metadata.txt", metadata_text)

        zip_buffer.seek(0)
        return zip_buffer.getvalue()

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Complete export error: {str(e)}")


def _format_complete_student_profile(profile_data: dict) -> str:
    """Format complete student profile as readable text"""
    lines = [
        "=" * 80,
        "COMPLETE STUDENT PROFILE",
        "=" * 80,
        "",
        "=== BASIC INFORMATION ===",
        f"Name: {profile_data.get('first_name', '')} {profile_data.get('last_name', '')}",
        f"Student ID: {profile_data.get('student_id', 'N/A')}",
        f"User ID: {profile_data.get('user_id', 'N/A')}",
        f"Email: {profile_data.get('email', 'N/A')}",
        f"GT Email: {profile_data.get('gt_email', 'N/A')}",
        f"Year: {profile_data.get('year_name', 'N/A')}",
        f"Group Type: {profile_data.get('group_type', 'N/A')}",
        f"Profile Picture: {profile_data.get('profile_picture_url', 'N/A')}",
        "",
        "=== ACHIEVEMENTS & POWERPOINT ===",
        f"PowerPoint Embed URL: {profile_data.get('ppt_embed_url', 'N/A')}",
        f"PowerPoint Edit URL: {profile_data.get('ppt_edit_url', 'N/A')}",
        "",
        "=== STRENGTHS ===",
    ]

    strengths = profile_data.get('strengths', [])
    if strengths:
        for strength in strengths:
            lines.append(f"  • {strength}")
    else:
        lines.append("  No strengths listed")

    lines.extend(["", "=== CHALLENGES ==="])
    challenges = profile_data.get('challenges', [])
    if challenges:
        for challenge in challenges:
            lines.append(f"  • {challenge}")
    else:
        lines.append("  No challenges listed")

    lines.extend([
        "",
        "=== GOALS ===",
        f"Long-term Goals: {profile_data.get('long_term_goals', 'N/A')}",
        f"Short-term Goals: {profile_data.get('short_term_goals', 'N/A')}",
        "",
        "=== BEST WAYS TO HELP ===",
    ])

    best_ways = profile_data.get('best_ways_to_help', [])
    if best_ways:
        for way in best_ways:
            lines.append(f"  • {way}")
    else:
        lines.append("  No best ways listed")

    lines.extend([
        "",
        "=== HOBBIES & INTERESTS ===",
        profile_data.get('hobbies_and_interests', 'N/A'),
        "",
    ])

    # Profile summaries if available
    summaries = profile_data.get('profile_summaries', {})
    if summaries:
        lines.extend([
            "=== AI-GENERATED SUMMARIES ===",
            "",
            "Strengths Summary:",
            summaries.get('strengths_short', 'N/A'),
            "",
            "Short-term Goals Summary:",
            summaries.get('short_term_goals', 'N/A'),
            "",
            "Long-term Goals Summary:",
            summaries.get('long_term_goals', 'N/A'),
            "",
            "Best Ways to Help Summary:",
            summaries.get('best_ways_to_help', 'N/A'),
            "",
            "Vision Statement:",
            summaries.get('vision', 'N/A'),
        ])

    lines.append("=" * 80)
    return "\n".join(lines)


def _create_all_students_summary_csv(students_data: List[dict]) -> str:
    """Create CSV summary of all students with their assignment counts"""
    output = io.StringIO()
    writer = csv.writer(output)

    # Header
    writer.writerow([
        "Student ID",
        "First Name",
        "Last Name",
        "Email",
        "Year",
        "Total Assignments",
        "Total Versions",
        "Has Profile"
    ])

    # Data rows
    for student in students_data:
        writer.writerow([
            student.get("student_id", ""),
            student.get("first_name", ""),
            student.get("last_name", ""),
            student.get("email", ""),
            student.get("year_name", ""),
            student.get("total_assignments", 0),
            student.get("total_versions", 0),
            "Yes" if student.get("has_profile") else "No"
        ])

    return output.getvalue()


def _format_export_metadata(assignment_data: dict, profile_data: dict) -> str:
    """Format export metadata"""
    from datetime import datetime as dt
    lines = [
        "=" * 60,
        "EXPORT METADATA",
        "=" * 60,
        f"Export Date: {dt.utcnow().isoformat()}",
        f"Student ID: {profile_data.get('student_id', 'N/A')}",
        f"Student Name: {profile_data.get('first_name', '')} {profile_data.get('last_name', '')}",
        f"Total Assignments: {assignment_data['export_metadata']['total_assignments']}",
        f"Total Classes: {len(assignment_data['classes'])}",
        "",
        "=== CONTENTS ===",
        "  • student_profile.txt - Complete student profile",
        "  • student_profile.json - Profile in JSON format",
        "  • classes_and_learning_goals.txt - Enrolled classes and goals",
        "  • assignments_summary.csv - Spreadsheet of all assignments",
        "  • assignments/ - Folder containing all assignments",
        "    ├── Each assignment has its own folder",
        "    ├── original_assignment.docx - Original assignment",
        "    ├── version_N.docx - Generated content (final version)",
        "    ├── version_N_complete_details.txt - Complete metadata including:",
        "    │   • Skills for success",
        "    │   • Learning pathways (generated options) with full reasoning",
        "    │   • Student's ideas for changes",
        "    │   • Selected options",
        "    │   • Generation history",
        "    │   • Rating history",
        "    └── ratings.txt - Current student feedback and ratings",
        "",
        "This export contains ALL data for this student including:",
        "  ✓ Complete profile with strengths, challenges, goals",
        "  ✓ All class enrollments with learning goals",
        "  ✓ All assignments with original content",
        "  ✓ All assignment versions and regenerations",
        "  ✓ Learning pathways with full reasoning for each version",
        "  ✓ Skills for success for each version",
        "  ✓ Student's ideas and additional suggestions",
        "  ✓ All ratings and feedback history",
        "  ✓ Rating history showing changes over time",
        "  ✓ Generation history showing content evolution",
        "  ✓ PowerPoint achievements tracking",
        "=" * 60
    ]
    return "\n".join(lines)


def export_all_students_complete_data(
    student_ids: Optional[List[int]] = None,
    assignment_ids: Optional[List[int]] = None
) -> bytes:
    """
    Export complete data for ALL students (or filtered subset) in one comprehensive ZIP.

    Creates a master ZIP file containing individual student folders, each with complete data.

    Args:
        student_ids: Optional list of student IDs to filter by (exports ALL if None)
        assignment_ids: Optional list of assignment IDs to filter by

    Returns:
        ZIP file as bytes with all student data
    """
    try:
        from application.features.students.crud import fetch_all_students_with_names
        from application.features.student_profile.crud import get_complete_profile

        # Get all students (or filtered list)
        all_students = fetch_all_students_with_names()

        # Filter by student_ids if provided
        if student_ids:
            all_students = [s for s in all_students if s.get('id') in student_ids]

        if not all_students:
            raise HTTPException(status_code=404, detail="No students found")

        # Create master ZIP
        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as master_zip:
            students_summary_data = []

            # Process each student
            for student in all_students:
                student_id = student.get('id')
                first_name = student.get('first_name', 'Unknown')
                last_name = student.get('last_name', 'Unknown')

                # Sanitize names for folder
                safe_first = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in first_name)
                safe_last = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in last_name)
                student_folder = f"student_{student_id}_{safe_first}_{safe_last}"

                try:
                    # Get student's complete export data
                    student_zip_bytes = export_complete_student_data(student_id, assignment_ids)

                    # Extract the student's ZIP and add contents to master ZIP under student folder
                    student_zip = zipfile.ZipFile(io.BytesIO(student_zip_bytes), 'r')

                    for file_info in student_zip.filelist:
                        file_data = student_zip.read(file_info.filename)
                        master_zip.writestr(f"{student_folder}/{file_info.filename}", file_data)

                    student_zip.close()

                    # Gather summary data
                    # Get assignment count for summary
                    assignment_data = export_student_assignments_json(student_id, assignment_ids)
                    total_assignments = len(assignment_data.get('assignments', []))
                    total_versions = sum(len(a.get('versions', [])) for a in assignment_data.get('assignments', []))

                    # Check if profile exists
                    has_profile = False
                    try:
                        profile = get_complete_profile(student_id)
                        has_profile = profile is not None
                    except:
                        has_profile = False

                    students_summary_data.append({
                        "student_id": student_id,
                        "first_name": first_name,
                        "last_name": last_name,
                        "email": student.get('email', 'N/A'),
                        "year_name": student.get('year_name', 'N/A'),
                        "total_assignments": total_assignments,
                        "total_versions": total_versions,
                        "has_profile": has_profile
                    })

                except Exception as e:
                    # Log error but continue with other students
                    error_msg = f"Error exporting student {student_id}: {str(e)}\n"
                    master_zip.writestr(f"{student_folder}/EXPORT_ERROR.txt", error_msg)

                    students_summary_data.append({
                        "student_id": student_id,
                        "first_name": first_name,
                        "last_name": last_name,
                        "email": student.get('email', 'N/A'),
                        "year_name": student.get('year_name', 'N/A'),
                        "total_assignments": 0,
                        "total_versions": 0,
                        "has_profile": False
                    })

            # Add summary CSV at root level
            summary_csv = _create_all_students_summary_csv(students_summary_data)
            master_zip.writestr("export_summary.csv", summary_csv)

            # Add master export metadata
            from datetime import datetime as dt
            master_metadata = [
                "=" * 80,
                "ALL STUDENTS EXPORT METADATA",
                "=" * 80,
                f"Export Date: {dt.utcnow().isoformat()}",
                f"Total Students Exported: {len(students_summary_data)}",
                f"Filtered by Student IDs: {student_ids if student_ids else 'No (all students)'}",
                f"Filtered by Assignment IDs: {assignment_ids if assignment_ids else 'No (all assignments)'}",
                "",
                "=== STRUCTURE ===",
                "This ZIP contains individual folders for each student with:",
                "  • Complete student profile (txt and json)",
                "  • All class enrollments with learning goals",
                "  • All assignments with original content",
                "  • All assignment versions with complete metadata",
                "  • Learning pathways with full reasoning",
                "  • Skills for success for each version",
                "  • Student's ideas and suggestions",
                "  • All ratings and feedback history",
                "  • Rating and generation history",
                "",
                "=== FILES ===",
                "  • export_summary.csv - Overview of all students",
                "  • student_{id}_{name}/ - Individual student folders",
                "",
                "See individual student folders for complete data.",
                "=" * 80
            ]
            master_zip.writestr("export_metadata.txt", "\n".join(master_metadata))

        zip_buffer.seek(0)
        return zip_buffer.getvalue()

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"All students export error: {str(e)}")
